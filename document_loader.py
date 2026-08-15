"""
document_loader.py
-------------------
Extraction de texte à partir de fichiers PDF, DOCX et TXT.

Ce module n'a aucune dépendance à Streamlit : il peut être testé et réutilisé
de façon totalement indépendante de l'interface.

Bonus : les pages de PDF scannées (images sans couche de texte) sont
automatiquement détectées et passées à l'OCR (PyMuPDF + pytesseract) pour
en extraire le texte. L'OCR échoue silencieusement (renvoie une chaîne
vide) si Tesseract n'est pas installé sur la machine, ce qui dégrade
proprement vers un simple avertissement plutôt qu'un plantage.
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import List, Optional, TypedDict, Union

from pypdf import PdfReader
from docx import Document as DocxDocument

PathLike = Union[str, Path]

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# En dessous de ce nombre de caractères extraits nativement, une page PDF est
# considérée comme "sans couche de texte" (probablement scannée) et bascule
# sur l'OCR.
MIN_CHARS_FOR_TEXT_LAYER = 20

# Langues essayées pour l'OCR, dans l'ordre (repli automatique si une langue
# n'est pas installée sur la machine hôte).
OCR_LANGUAGE_CANDIDATES = ("fra+eng", "eng", None)


class UnsupportedFileType(ValueError):
    """Levée quand un type de fichier non pris en charge est fourni."""


class TextSegment(TypedDict):
    """Un segment de texte extrait d'un document.

    `page` vaut None pour les formats sans notion de page (DOCX, TXT).
    `ocr` indique si le texte de ce segment provient de l'OCR plutôt que
    d'une extraction de texte native.
    """

    page: Optional[int]
    text: str
    ocr: bool


def _ocr_pdf_page(file_path: PathLike, page_index: int, dpi: int = 200) -> str:
    """OCRise une page de PDF (rendue en image) via PyMuPDF + pytesseract.

    Renvoie une chaîne vide en cas d'échec : dépendance absente (PyMuPDF /
    pytesseract / Pillow) ou binaire Tesseract non installé sur la machine.
    Cela permet à l'extraction PDF de dégrader proprement plutôt que de
    planter quand l'OCR n'est pas disponible.
    """
    try:
        import pymupdf
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""

    try:
        pdf_doc = pymupdf.open(str(file_path))
        try:
            page = pdf_doc[page_index]
            pixmap = page.get_pixmap(dpi=dpi)
            image = Image.open(io.BytesIO(pixmap.tobytes("png")))
        finally:
            pdf_doc.close()
    except Exception:
        return ""

    for lang in OCR_LANGUAGE_CANDIDATES:
        try:
            kwargs = {"lang": lang} if lang else {}
            return pytesseract.image_to_string(image, **kwargs)
        except Exception:
            continue
    return ""


def extract_pdf_pages(
    file_path: PathLike, ocr_fallback: bool = True
) -> List[dict]:
    """Renvoie une liste de {"text": str, "ocr": bool}, un élément par page.

    Si `ocr_fallback` est vrai (par défaut), toute page dont l'extraction de
    texte native est quasi vide (PDF scanné) est automatiquement passée à
    l'OCR.
    """
    reader = PdfReader(str(file_path))
    pages = []
    for index, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""

        used_ocr = False
        if ocr_fallback and len(text.strip()) < MIN_CHARS_FOR_TEXT_LAYER:
            ocr_text = _ocr_pdf_page(file_path, index)
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
                used_ocr = True

        pages.append({"text": text, "ocr": used_ocr})
    return pages


def extract_docx_text(file_path: PathLike) -> str:
    """Extrait le texte d'un document Word (.docx), paragraphes et tableaux inclus."""
    doc = DocxDocument(str(file_path))
    parts: List[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


def extract_txt_text(file_path: PathLike) -> str:
    """Lit un fichier texte brut en tentant plusieurs encodages courants."""
    raw_bytes = Path(file_path).read_bytes()
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw_bytes.decode("utf-8", errors="replace")


def load_document(file_path: PathLike, ocr_fallback: bool = True) -> List[TextSegment]:
    """Charge un document et renvoie une liste de segments de texte.

    - Pour un PDF : un segment par page (les pages vides sont ignorées).
      Les pages sans couche de texte (PDF scannés) sont automatiquement
      passées à l'OCR si `ocr_fallback` est vrai.
    - Pour un DOCX ou un TXT : un segment unique (pas de notion de page).

    Lève `UnsupportedFileType` si l'extension n'est pas prise en charge.
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        pages = extract_pdf_pages(file_path, ocr_fallback=ocr_fallback)
        return [
            {"page": index + 1, "text": page["text"], "ocr": page["ocr"]}
            for index, page in enumerate(pages)
            if page["text"] and page["text"].strip()
        ]

    if ext == ".docx":
        text = extract_docx_text(file_path)
        return [{"page": None, "text": text, "ocr": False}] if text.strip() else []

    if ext == ".txt":
        text = extract_txt_text(file_path)
        return [{"page": None, "text": text, "ocr": False}] if text.strip() else []

    raise UnsupportedFileType(
        f"Type de fichier non pris en charge : '{ext}'. "
        f"Formats acceptés : {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
    )


def full_text(segments: List[TextSegment]) -> str:
    """Concatène tous les segments d'un document en un seul texte."""
    return "\n\n".join(segment["text"] for segment in segments)


def ocr_page_count(segments: List[TextSegment]) -> int:
    """Nombre de segments dont le texte provient de l'OCR (utile pour
    informer l'utilisateur qu'un document scanné a été traité)."""
    return sum(1 for segment in segments if segment.get("ocr"))
