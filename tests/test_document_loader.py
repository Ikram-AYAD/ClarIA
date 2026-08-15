"""
tests/test_document_loader.py
------------------------------
Tests unitaires pour document_loader.py : extraction PDF / DOCX / TXT et
repli OCR sur les PDF scannes.

Tous les fichiers de test sont generes a la volee dans le repertoire
temporaire fourni par pytest (`tmp_path`) : aucun fichier binaire n'est
commite dans le depot.
"""

import shutil
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import document_loader as dl  # noqa: E402

TESSERACT_AVAILABLE = shutil.which("tesseract") is not None


def _make_pdf(path: Path, page_texts: list[str]) -> None:
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate

    styles = getSampleStyleSheet()
    story = []
    for i, text in enumerate(page_texts):
        if i > 0:
            story.append(PageBreak())
        story.append(Paragraph(text, styles["Normal"]))
    SimpleDocTemplate(str(path)).build(story)


def _make_docx(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Politique de teletravail", level=1)
    doc.add_paragraph("Les employes peuvent teletravailler jusqu'a trois jours par semaine.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Jour"
    table.rows[0].cells[1].text = "Statut"
    table.rows[1].cells[0].text = "Lundi"
    table.rows[1].cells[1].text = "Teletravail autorise"
    doc.save(str(path))


def _make_scanned_pdf(path: Path, text: str) -> None:
    """Cree un PDF 'scanne' : une image contenant du texte, inseree dans un
    PDF sans aucune couche de texte native (donc non extractible par pypdf,
    seulement par OCR)."""
    from PIL import Image, ImageDraw
    import pymupdf

    image = Image.new("RGB", (900, 300), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 100), text, fill="black")
    image_path = path.with_suffix(".png")
    image.save(image_path)

    pdf_doc = pymupdf.open()
    page = pdf_doc.new_page(width=900, height=300)
    page.insert_image(page.rect, filename=str(image_path))
    pdf_doc.save(str(path))
    pdf_doc.close()
    image_path.unlink()


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def test_extract_pdf_pages_returns_one_entry_per_page(tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    _make_pdf(pdf_path, ["Premiere page. Contenu A.", "Deuxieme page. Contenu B."])

    pages = dl.extract_pdf_pages(pdf_path)
    assert len(pages) == 2
    assert "Contenu A" in pages[0]["text"]
    assert "Contenu B" in pages[1]["text"]
    assert pages[0]["ocr"] is False


def test_load_document_pdf_assigns_page_numbers(tmp_path):
    pdf_path = tmp_path / "sample.pdf"
    _make_pdf(pdf_path, ["Le budget est de 4.2 millions d'euros.", "Les risques sont limites."])

    segments = dl.load_document(pdf_path)
    assert len(segments) == 2
    assert segments[0]["page"] == 1
    assert segments[1]["page"] == 2
    assert "budget" in segments[0]["text"].lower()


@pytest.mark.skipif(not TESSERACT_AVAILABLE, reason="Tesseract OCR n'est pas installe sur cette machine")
def test_ocr_fallback_extracts_text_from_scanned_pdf(tmp_path):
    scanned_path = tmp_path / "scanned.pdf"
    _make_scanned_pdf(scanned_path, "INVOICE TOTAL 4200")

    segments = dl.load_document(scanned_path, ocr_fallback=True)
    assert len(segments) == 1
    assert segments[0]["ocr"] is True
    # L'OCR n'est pas parfait a 100%, mais doit retrouver l'essentiel du texte.
    assert "4200" in segments[0]["text"] or "INVOICE" in segments[0]["text"].upper()
    assert dl.ocr_page_count(segments) == 1


def test_ocr_disabled_leaves_scanned_pdf_empty(tmp_path):
    scanned_path = tmp_path / "scanned.pdf"
    _make_scanned_pdf(scanned_path, "INVOICE TOTAL 4200")

    segments = dl.load_document(scanned_path, ocr_fallback=False)
    # Sans OCR, aucune couche de texte n'est trouvee : aucun segment exploitable.
    assert segments == []


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def test_extract_docx_text_includes_paragraphs_and_tables(tmp_path):
    docx_path = tmp_path / "sample.docx"
    _make_docx(docx_path)

    text = dl.extract_docx_text(docx_path)
    assert "teletravail" in text.lower()
    assert "Jour" in text and "Statut" in text


def test_load_document_docx_has_no_page_number(tmp_path):
    docx_path = tmp_path / "sample.docx"
    _make_docx(docx_path)

    segments = dl.load_document(docx_path)
    assert len(segments) == 1
    assert segments[0]["page"] is None
    assert segments[0]["ocr"] is False


# --------------------------------------------------------------------------
# TXT
# --------------------------------------------------------------------------

def test_load_document_txt_reads_utf8(tmp_path):
    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("Reunion du 12 aout : decision reportee.", encoding="utf-8")

    segments = dl.load_document(txt_path)
    assert len(segments) == 1
    assert "Reunion" in segments[0]["text"]
    assert segments[0]["page"] is None


# --------------------------------------------------------------------------
# Types non pris en charge
# --------------------------------------------------------------------------

def test_load_document_unsupported_extension_raises(tmp_path):
    bad_path = tmp_path / "notes.xyz"
    bad_path.write_text("contenu", encoding="utf-8")

    with pytest.raises(dl.UnsupportedFileType):
        dl.load_document(bad_path)


def test_full_text_concatenates_segments():
    segments = [
        {"page": 1, "text": "Premiere partie.", "ocr": False},
        {"page": 2, "text": "Deuxieme partie.", "ocr": False},
    ]
    combined = dl.full_text(segments)
    assert "Premiere partie." in combined
    assert "Deuxieme partie." in combined
